"""
文档处理模块

负责文档的加载、解析、切块等预处理操作，是GraphRAG管道的起始阶段。
"""

import asyncio
import hashlib
import pandas as pd
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Iterator, AsyncIterator
from dataclasses import dataclass, field
import tiktoken
from abc import ABC, abstractmethod
import logging

# 导入文档解析库
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    from bs4 import BeautifulSoup
    HAS_BEAUTIFULSOUP = True
except ImportError:
    HAS_BEAUTIFULSOUP = False

from ..config.settings import DocumentProcessingConfig
from ..utils.exceptions import DocumentProcessingError, ValidationError

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """文档块数据类
    
    表示文档切分后的单个块。
    """
    
    chunk_id: str  # 块的唯一标识符
    content: str  # 块的文本内容
    token_count: int  # token数量
    start_char: int  # 在原文档中的起始字符位置
    end_char: int  # 在原文档中的结束字符位置
    source_document: str  # 源文档路径
    metadata: Dict[str, Any] = field(default_factory=dict)  # 额外元数据
    
    def __post_init__(self):
        """初始化后处理"""
        if not self.chunk_id:
            # 生成基于内容的哈希ID
            content_hash = hashlib.md5(self.content.encode()).hexdigest()[:8]
            self.chunk_id = f"chunk_{content_hash}"


@dataclass
class Document:
    """文档数据类
    
    表示一个完整的源文档。
    """
    
    document_id: str  # 文档的唯一标识符
    content: str  # 文档的完整文本内容
    file_path: Path  # 文档文件路径
    file_type: str  # 文档类型
    file_size: int  # 文件大小(字节)
    created_at: Optional[str] = None  # 创建时间
    modified_at: Optional[str] = None  # 修改时间
    metadata: Dict[str, Any] = field(default_factory=dict)  # 文档元数据
    
    def __post_init__(self):
        """初始化后处理"""
        if not self.document_id:
            # 生成基于文件路径的哈希ID
            path_hash = hashlib.md5(str(self.file_path).encode()).hexdigest()[:8]
            self.document_id = f"doc_{path_hash}"


class BaseDocumentReader(ABC):
    """文档读取器基类
    
    定义所有文档读取器必须实现的接口。
    """
    
    @abstractmethod
    def can_handle(self, file_path: Path) -> bool:
        """判断是否能处理指定的文件类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            是否能处理
        """
        pass
    
    @abstractmethod
    def read_document(self, file_path: Path) -> str:
        """读取文档内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档文本内容
            
        Raises:
            DocumentProcessingError: 当读取失败时
        """
        pass


class TextDocumentReader(BaseDocumentReader):
    """纯文本文档读取器
    
    处理.txt、.md等纯文本文件。
    """
    
    def can_handle(self, file_path: Path) -> bool:
        """判断是否能处理指定的文件类型"""
        text_extensions = {'.txt', '.md', '.markdown', '.rst', '.log'}
        return file_path.suffix.lower() in text_extensions
    
    def read_document(self, file_path: Path) -> str:
        """读取纯文本文档"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # 所有编码都失败
            raise DocumentProcessingError(
                f"Unable to decode file with any supported encoding: {file_path}",
                document_path=str(file_path)
            )
            
        except FileNotFoundError:
            raise DocumentProcessingError(
                f"File not found: {file_path}",
                document_path=str(file_path)
            )
        except Exception as e:
            raise DocumentProcessingError(
                f"Error reading text document: {e}",
                document_path=str(file_path)
            )


class PDFDocumentReader(BaseDocumentReader):
    """PDF文档读取器
    
    使用PyMuPDF处理PDF文件。
    """
    
    def can_handle(self, file_path: Path) -> bool:
        """判断是否能处理指定的文件类型"""
        return file_path.suffix.lower() == '.pdf' and HAS_PYMUPDF
    
    def read_document(self, file_path: Path) -> str:
        """读取PDF文档"""
        if not HAS_PYMUPDF:
            raise DocumentProcessingError(
                "PyMuPDF not installed. Please install with: pip install pymupdf",
                document_path=str(file_path)
            )
        
        try:
            doc = fitz.open(file_path)
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():  # 只添加非空页面
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
            
            doc.close()
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Error reading PDF document: {e}",
                document_path=str(file_path)
            )


class DocxDocumentReader(BaseDocumentReader):
    """Word文档读取器
    
    使用python-docx处理.docx文件。
    """
    
    def can_handle(self, file_path: Path) -> bool:
        """判断是否能处理指定的文件类型"""
        return file_path.suffix.lower() == '.docx' and HAS_PYTHON_DOCX
    
    def read_document(self, file_path: Path) -> str:
        """读取Word文档"""
        if not HAS_PYTHON_DOCX:
            raise DocumentProcessingError(
                "python-docx not installed. Please install with: pip install python-docx",
                document_path=str(file_path)
            )
        
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Error reading DOCX document: {e}",
                document_path=str(file_path)
            )


class HTMLDocumentReader(BaseDocumentReader):
    """HTML文档读取器
    
    使用BeautifulSoup处理HTML文件。
    """
    
    def can_handle(self, file_path: Path) -> bool:
        """判断是否能处理指定的文件类型"""
        html_extensions = {'.html', '.htm', '.xhtml'}
        return file_path.suffix.lower() in html_extensions and HAS_BEAUTIFULSOUP
    
    def read_document(self, file_path: Path) -> str:
        """读取HTML文档"""
        if not HAS_BEAUTIFULSOUP:
            raise DocumentProcessingError(
                "BeautifulSoup not installed. Please install with: pip install beautifulsoup4",
                document_path=str(file_path)
            )
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # 移除脚本和样式标签
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 提取文本内容
            text = soup.get_text()
            
            # 清理空白字符
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Error reading HTML document: {e}",
                document_path=str(file_path)
            )


class DocumentReaderFactory:
    """文档读取器工厂类
    
    根据文件类型选择合适的读取器。
    """
    
    def __init__(self):
        """初始化工厂"""
        self._readers = [
            TextDocumentReader(),
            PDFDocumentReader(),
            DocxDocumentReader(),
            HTMLDocumentReader(),
        ]
    
    def get_reader(self, file_path: Path) -> BaseDocumentReader:
        """获取适合的文档读取器
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档读取器实例
            
        Raises:
            DocumentProcessingError: 当找不到合适的读取器时
        """
        for reader in self._readers:
            if reader.can_handle(file_path):
                return reader
        
        # 没有找到合适的读取器
        supported_formats = []
        for reader in self._readers:
            if hasattr(reader, 'supported_formats'):
                supported_formats.extend(reader.supported_formats)
        
        raise DocumentProcessingError(
            f"Unsupported file format: {file_path.suffix}. "
            f"Supported formats: {supported_formats}",
            document_path=str(file_path)
        )
    
    def register_reader(self, reader: BaseDocumentReader):
        """注册新的文档读取器
        
        Args:
            reader: 文档读取器实例
        """
        if not isinstance(reader, BaseDocumentReader):
            raise ValidationError("Reader must inherit from BaseDocumentReader")
        
        self._readers.append(reader)


class DocumentChunker:
    """文档切块器
    
    负责将长文档切分为适合LLM处理的较小块。
    """
    
    def __init__(self, config: DocumentProcessingConfig):
        """初始化切块器
        
        Args:
            config: 文档处理配置
        """
        self.config = config
        
        # 初始化token编码器
        try:
            self._encoding = tiktoken.get_encoding(config.encoding_name)
        except Exception:
            logger.warning(f"Unable to load encoding {config.encoding_name}, using cl100k_base")
            self._encoding = tiktoken.get_encoding("cl100k_base")
    
    def chunk_document(self, document: Document) -> List[DocumentChunk]:
        """对文档进行切块
        
        Args:
            document: 待切块的文档
            
        Returns:
            文档块列表
        """
        text = document.content
        chunks = []
        
        # 按句子分割文本
        sentences = self._split_into_sentences(text)
        
        current_chunk = []
        current_tokens = 0
        start_char = 0
        
        for sentence in sentences:
            sentence_tokens = self._count_tokens(sentence)
            
            # 检查是否需要开始新的块
            if (current_tokens + sentence_tokens > self.config.chunk_size and 
                current_chunk):
                
                # 创建当前块
                chunk_content = ' '.join(current_chunk)
                end_char = start_char + len(chunk_content)
                
                chunk = DocumentChunk(
                    chunk_id="",  # 将在__post_init__中生成
                    content=chunk_content,
                    token_count=current_tokens,
                    start_char=start_char,
                    end_char=end_char,
                    source_document=str(document.file_path),
                    metadata={
                        "document_id": document.document_id,
                        "chunk_index": len(chunks),
                        "total_chunks": -1  # 稍后更新
                    }
                )
                chunks.append(chunk)
                
                # 处理重叠
                overlap_content = self._create_overlap(current_chunk, current_tokens)
                current_chunk = overlap_content
                current_tokens = sum(self._count_tokens(s) for s in current_chunk)
                
                # 更新起始位置（考虑重叠）
                if overlap_content:
                    overlap_text = ' '.join(overlap_content)
                    start_char = end_char - len(overlap_text)
                else:
                    start_char = end_char
            
            current_chunk.append(sentence)
            current_tokens += sentence_tokens
        
        # 处理最后一个块
        if current_chunk:
            chunk_content = ' '.join(current_chunk)
            end_char = start_char + len(chunk_content)
            
            chunk = DocumentChunk(
                chunk_id="",
                content=chunk_content,
                token_count=current_tokens,
                start_char=start_char,
                end_char=end_char,
                source_document=str(document.file_path),
                metadata={
                    "document_id": document.document_id,
                    "chunk_index": len(chunks),
                    "total_chunks": len(chunks) + 1
                }
            )
            chunks.append(chunk)
        
        # 更新所有块的total_chunks
        for chunk in chunks:
            chunk.metadata["total_chunks"] = len(chunks)
        
        logger.info(f"Document {document.document_id} chunked into {len(chunks)} chunks")
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """将文本分割为句子
        
        Args:
            text: 输入文本
            
        Returns:
            句子列表
        """
        # 简单的句子分割实现
        # 在实际应用中可能需要更复杂的NLP工具
        import re
        
        # 按句号、感叹号、问号分割
        sentence_endings = re.compile(r'[.!?]+\s+')
        sentences = sentence_endings.split(text)
        
        # 清理空句子
        sentences = [s.strip() for s in sentences if s.strip()]
        
        return sentences
    
    def _count_tokens(self, text: str) -> int:
        """计算文本的token数量
        
        Args:
            text: 输入文本
            
        Returns:
            token数量
        """
        try:
            return len(self._encoding.encode(text))
        except Exception:
            # 如果编码失败，使用粗略估计
            return len(text) // 4
    
    def _create_overlap(self, current_chunk: List[str], current_tokens: int) -> List[str]:
        """创建块之间的重叠内容
        
        Args:
            current_chunk: 当前块的句子列表
            current_tokens: 当前块的token数量
            
        Returns:
            重叠部分的句子列表
        """
        if not current_chunk or self.config.chunk_overlap <= 0:
            return []
        
        overlap_sentences = []
        overlap_tokens = 0
        
        # 从后往前添加句子，直到达到重叠大小
        for sentence in reversed(current_chunk):
            sentence_tokens = self._count_tokens(sentence)
            
            if overlap_tokens + sentence_tokens <= self.config.chunk_overlap:
                overlap_sentences.insert(0, sentence)
                overlap_tokens += sentence_tokens
            else:
                break
        
        return overlap_sentences


class DocumentProcessor:
    """文档处理器主类
    
    协调文档读取、解析和切块的整个流程。
    """
    
    def __init__(self, config: DocumentProcessingConfig):
        """初始化文档处理器
        
        Args:
            config: 文档处理配置
        """
        self.config = config
        self.reader_factory = DocumentReaderFactory()
        self.chunker = DocumentChunker(config)
    
    def process_file(self, file_path: Union[str, Path]) -> List[DocumentChunk]:
        """处理单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档块列表
            
        Raises:
            DocumentProcessingError: 当处理失败时
        """
        file_path = Path(file_path)
        
        # 验证文件存在
        if not file_path.exists():
            raise DocumentProcessingError(
                f"File does not exist: {file_path}",
                document_path=str(file_path)
            )
        
        # 验证文件大小
        file_size = file_path.stat().st_size
        max_size = 50 * 1024 * 1024  # 50MB限制
        if file_size > max_size:
            raise DocumentProcessingError(
                f"File too large: {file_size} bytes (max: {max_size})",
                document_path=str(file_path)
            )
        
        try:
            # 获取合适的读取器
            reader = self.reader_factory.get_reader(file_path)
            
            # 读取文档内容
            content = reader.read_document(file_path)
            
            # 创建文档对象
            document = Document(
                document_id="",  # 将在__post_init__中生成
                content=content,
                file_path=file_path,
                file_type=file_path.suffix.lower(),
                file_size=file_size,
                metadata={
                    "processed_at": str(pd.Timestamp.now()),
                    "processor_version": "1.0.0"
                }
            )
            
            # 切块处理
            chunks = self.chunker.chunk_document(document)
            
            logger.info(f"Successfully processed {file_path}: {len(chunks)} chunks")
            return chunks
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                f"Unexpected error processing file: {e}",
                document_path=str(file_path)
            )
    
    def process_directory(
        self, 
        directory_path: Union[str, Path],
        recursive: bool = True
    ) -> Dict[str, List[DocumentChunk]]:
        """处理目录中的所有文件
        
        Args:
            directory_path: 目录路径
            recursive: 是否递归处理子目录
            
        Returns:
            文件路径到文档块列表的映射
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise DocumentProcessingError(
                f"Directory does not exist: {directory_path}",
                document_path=str(directory_path)
            )
        
        results = {}
        
        # 获取所有支持的文件
        pattern = "**/*" if recursive else "*"
        for file_path in directory_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.config.supported_formats:
                try:
                    chunks = self.process_file(file_path)
                    results[str(file_path)] = chunks
                except DocumentProcessingError as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    # 继续处理其他文件
                    continue
        
        logger.info(f"Processed {len(results)} files from {directory_path}")
        return results
    
    async def process_files_async(
        self, 
        file_paths: List[Union[str, Path]],
        max_concurrent: int = 5
    ) -> Dict[str, List[DocumentChunk]]:
        """异步处理多个文件
        
        Args:
            file_paths: 文件路径列表
            max_concurrent: 最大并发数
            
        Returns:
            文件路径到文档块列表的映射
        """
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def _process_file_async(file_path: Union[str, Path]) -> tuple:
            """异步处理单个文件"""
            async with semaphore:
                try:
                    # 在线程池中执行同步处理
                    loop = asyncio.get_event_loop()
                    chunks = await loop.run_in_executor(
                        None, self.process_file, file_path
                    )
                    return str(file_path), chunks
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    return str(file_path), None
        
        # 并发处理所有文件
        tasks = [_process_file_async(fp) for fp in file_paths]
        results = await asyncio.gather(*tasks)
        
        # 过滤成功的结果
        return {path: chunks for path, chunks in results if chunks is not None}